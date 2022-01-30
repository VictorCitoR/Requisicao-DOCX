import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import PySimpleGUI as SG


# Define função para marcação de caixas
def add_caixa(run, sit, t):
    if sit is False:
        run.add_picture('vazio.png', width=Cm(0.33), height=Cm(0.33))
    else:
        run.add_picture('marcado.png', width=Cm(0.33), height=Cm(0.33))
    run.add_text(f' {t}')


# Criação das janelas
def janela_nova(j, pset=None):

    if True:
        frame_ano = [
            [SG.Input(size=(4, 0), key=f'{j}: ano')]
        ]
        frame_requisicao = [
            [SG.Input(size=(10, 0), key=f'{j}: requisicao')]
        ]
        frame_coordenador = [
            [SG.Input(size=(20, 0), key=f'{j}: coordenador')]
        ]
        frame_DS = [
            [SG.Input(size=(10, 0), key=f'{j}: DS')]
        ]
        frame_professores = [
            [SG.Input(key=f'{j}: professores')]
        ]
        frame_supervisor = [
            [SG.Input(key=f'{j}: supervisor')]
        ]
        frame_sede = [
            [SG.T('Aldeota      Centro')],
            [SG.CB('EBS', key=f'{j}: EBS'), SG.CB('EGS', key=f'{j}: EGS'), SG.CB('NGS', key=f'{j}: NGS')]
        ]
        frame_turno = [
            [SG.CB('M', key=f'{j}: M'), SG.CB('T', key=f'{j}: T'), SG.CB('N', key=f'{j}: N')]
        ]
        frame_etapa = [
            [SG.CB('1ª     ', key=f'{j}: 1et'), SG.CB('4ª  ', key=f'{j}: 4et')],
            [SG.CB('2ª     ', key=f'{j}: 2et'), SG.CB('2ªVG', key=f'{j}: 2vg')],
            [SG.CB('1ªVG', key=f'{j}: 1vg'), SG.CB('Rec.', key=f'{j}: rec')],
            [SG.CB('3ª     ', key=f'{j}: 3et'), SG.CB('AF  ', key=f'{j}: aff')]
        ]
        frame_ensino = [
            [SG.CB('Infantil', key=f'{j}: infan')],
            [SG.CB('Fund. I', key=f'{j}: fund1')],
            [SG.CB('Fund. II', key=f'{j}: fund2')],
            [SG.CB('Médio', key=f'{j}: medio')]
        ]
        frame_serie = [
            [SG.CB('1º', key=f'{j}: 1ano'), SG.CB('2º', key=f'{j}: 2ano'), SG.CB('3º', key=f'{j}: 3ano'),
             SG.CB('4º', key=f'{j}: 4ano'), SG.CB('5º', key=f'{j}: 5ano'), SG.CB('6º', key=f'{j}: 6ano'),
             SG.CB('7º', key=f'{j}: 7ano')],
            [SG.CB('8º', key=f'{j}: 8ano'), SG.CB('9º', key=f'{j}: 9ano'), SG.CB('Ext.', key=f'{j}: ext.'),
             SG.CB('Outro', key=f'{j}: out'),
             SG.Input(size=(20, 0), key=f'{j}: serie')]
        ]
        frame_turmas = [
            [SG.Input(size=(53, 0), key=f'{j}: turmas')]
        ]
        frame_natureza = [
            [SG.CB('NP', key=f'{j}: NP'), SG.CB('VG', key=f'{j}: VG'), SG.CB('Miniteste', key=f'{j}: MT')],
            [SG.CB('AP', key=f'{j}: AP'), SG.CB('Simulados', key=f'{j}: SIM'),
             SG.CB('SIMULADO ENEM', key=f'{j}: ENEM')],
            [SG.CB('TD/TODAS', key=f'{j}: TD'), SG.CB('Pautas de Avaliação', key=f'{j}: PA')]
        ]
        frame_observacoes = [
            [SG.Input(size=(110, 0), key=f'{j}: observacoes')]
        ]
        frame_envio = [
            [SG.Input(size=(10, 0), key=f'{j}: envio')]
        ]
        frame_aplicacao = [
            [SG.Input(size=(10, 0), key=f'{j}: aplicacao')]
        ]
        frame_SOP = [
            [SG.CB('Sim', key=f'{j}: Sim'), SG.CB('Não', key=f'{j}: Nao'),
             SG.Text('Questões:'), SG.Input(size=(54, 0), key=f'{j}: SOP')]
        ]

    valores_preset = {'presetName': ''}
    if pset is not None:
        preset_selecionada_na_funcao = open(f'{caminho}\\presets\\{pset}', 'r')
        texto_da_preset = preset_selecionada_na_funcao.read()
        preset_selecionada_na_funcao.close()
        linhas = texto_da_preset.split('\n')
        for linha_das_linhas in linhas:
            valores_da_linha = linha_das_linhas.split(': ')
            valores_preset[valores_da_linha[0]] = valores_da_linha[1]
        frame_ano = [
            [SG.Input(valores_preset['ano'], size=(4, 0), key=f'{j}: ano')]
        ]
        frame_requisicao = [
            [SG.Input(valores_preset['requisicao'], size=(10, 0), key=f'{j}: requisicao')]
        ]
        frame_coordenador = [
            [SG.Input(valores_preset['coordenador'], size=(20, 0), key=f'{j}: coordenador')]
        ]
        frame_DS = [
            [SG.Input(valores_preset['DS'], size=(10, 0), key=f'{j}: DS')]
        ]
        frame_professores = [
            [SG.Input(valores_preset['professores'], key=f'{j}: professores')]
        ]
        frame_supervisor = [
            [SG.Input(valores_preset['supervisor'], key=f'{j}: supervisor')]
        ]
        frame_sede = [
            [SG.T('Aldeota      Centro')],
            [SG.CB('EBS', key=f'{j}: EBS', default=eval(valores_preset['EBS'])),
             SG.CB('EGS', key=f'{j}: EGS', default=eval(valores_preset['EGS'])),
             SG.CB('NGS', key=f'{j}: NGS', default=eval(valores_preset['NGS']))]
        ]
        frame_turno = [
            [SG.CB('M', key=f'{j}: M', default=eval(valores_preset['M'])),
             SG.CB('T', key=f'{j}: T', default=eval(valores_preset['T'])),
             SG.CB('N', key=f'{j}: N', default=eval(valores_preset['N']))]
        ]
        frame_etapa = [
            [SG.CB('1ª     ', key=f'{j}: 1et', default=eval(valores_preset['1et'])),
             SG.CB('4ª  ', key=f'{j}: 4et', default=eval(valores_preset['4et']))],
            [SG.CB('2ª     ', key=f'{j}: 2et', default=eval(valores_preset['2et'])),
             SG.CB('2ªVG', key=f'{j}: 2vg', default=eval(valores_preset['2vg']))],
            [SG.CB('1ªVG', key=f'{j}: 1vg', default=eval(valores_preset['1vg'])),
             SG.CB('Rec.', key=f'{j}: rec', default=eval(valores_preset['rec']))],
            [SG.CB('3ª     ', key=f'{j}: 3et', default=eval(valores_preset['3et'])),
             SG.CB('AF  ', key=f'{j}: aff', default=eval(valores_preset['aff']))]
        ]
        frame_ensino = [
            [SG.CB('Infantil', key=f'{j}: infan', default=eval(valores_preset['infan']))],
            [SG.CB('Fund. I', key=f'{j}: fund1', default=eval(valores_preset['fund1']))],
            [SG.CB('Fund. II', key=f'{j}: fund2', default=eval(valores_preset['fund2']))],
            [SG.CB('Médio', key=f'{j}: medio', default=eval(valores_preset['medio']))]
        ]
        frame_serie = [
            [SG.CB('1º', key=f'{j}: 1ano', default=eval(valores_preset['1ano'])),
             SG.CB('2º', key=f'{j}: 2ano', default=eval(valores_preset['2ano'])),
             SG.CB('3º', key=f'{j}: 3ano', default=eval(valores_preset['3ano'])),
             SG.CB('4º', key=f'{j}: 4ano', default=eval(valores_preset['4ano'])),
             SG.CB('5º', key=f'{j}: 5ano', default=eval(valores_preset['5ano'])),
             SG.CB('6º', key=f'{j}: 6ano', default=eval(valores_preset['6ano'])),
             SG.CB('7º', key=f'{j}: 7ano', default=eval(valores_preset['7ano']))],
            [SG.CB('8º', key=f'{j}: 8ano', default=eval(valores_preset['8ano'])),
             SG.CB('9º', key=f'{j}: 9ano', default=eval(valores_preset['9ano'])),
             SG.CB('Ext.', key=f'{j}: ext.', default=eval(valores_preset['ext.'])),
             SG.CB('Outro', key=f'{j}: out', default=eval(valores_preset['out'])),
             SG.Input(valores_preset['serie'], size=(20, 0), key=f'{j}: serie')]
        ]
        frame_turmas = [
            [SG.Input(valores_preset['turmas'], size=(53, 0), key=f'{j}: turmas')]
        ]
        frame_natureza = [
            [SG.CB('NP', key=f'{j}: NP', default=eval(valores_preset['NP'])),
             SG.CB('VG', key=f'{j}: VG', default=eval(valores_preset['VG'])),
             SG.CB('Miniteste', key=f'{j}: MT', default=eval(valores_preset['MT']))],
            [SG.CB('AP', key=f'{j}: AP', default=eval(valores_preset['AP'])),
             SG.CB('Simulados', key=f'{j}: SIM', default=eval(valores_preset['SIM'])),
             SG.CB('SIMULADO ENEM', key=f'{j}: ENEM', default=eval(valores_preset['ENEM']))],
            [SG.CB('TD/TODAS', key=f'{j}: TD', default=eval(valores_preset['TD'])),
             SG.CB('Pautas de Avaliação', key=f'{j}: PA', default=eval(valores_preset['PA']))]
        ]
        frame_observacoes = [
            [SG.Input(valores_preset['observacoes'], size=(110, 0), key=f'{j}: observacoes')]
        ]
        frame_envio = [
            [SG.Input(valores_preset['envio'], size=(10, 0), key=f'{j}: envio')]
        ]
        frame_aplicacao = [
            [SG.Input(valores_preset['aplicacao'], size=(10, 0), key=f'{j}: aplicacao')]
        ]
        frame_SOP = [
                [SG.CB('Sim', key=f'{j}: Sim', default=eval(valores_preset['Sim'])),
                 SG.CB('Não', key=f'{j}: Nao', default=eval(valores_preset['Nao'])),
                 SG.Text('Questões:'), SG.Input(valores_preset['SOP'], size=(54, 0), key=f'{j}: SOP')]
            ]

    # layout
    layout = [
        [SG.Text('Colégio 7 de Setembro'), SG.Text('REQUISIÇÃO DE TRABALHO'),
         SG.Frame('Ano', frame_ano), SG.Frame('Requisição: ', frame_requisicao),
         SG.Text(f'{j}ª Requisição')],  # Cabeçalho
        [SG.Frame('Coordenador:', frame_coordenador), SG.Frame('Disciplina/Setor:', frame_DS),
         SG.Frame('Professores:', frame_professores)],  # Professores 1
        [SG.Frame('Supervisor:', frame_supervisor), SG.Frame('Sede:', frame_sede)],  # Professores 2
        [SG.Frame('Turno:', frame_turno), SG.Frame('Etapa:', frame_etapa), SG.Frame('Ensino:', frame_ensino),
         SG.Frame('Ano:', frame_serie)],  # Turno
        [SG.Frame('Turma(s):', frame_turmas), SG.Frame('Natureza do Trabalho:', frame_natureza)],  # Turmas
        [SG.Frame('Observações:', frame_observacoes)],  # Observação
        [SG.Frame('Envio:', frame_envio), SG.Frame('Aplicação:', frame_aplicacao),
         SG.Frame('SOP Correção:', frame_SOP)],  # Envio
        [SG.Button('Gerar Arquivo'), SG.Button('Salvar Preset'),
         SG.Input(valores_preset['presetName'], size=(10, 0), key='presetName'),
         SG.Button('Selecionar Preset'), SG.Button('Nova requisição')]
    ]
    # Janela
    return SG.Window('Requisicao', layout=layout, finalize=True)


def janela_preset(ps):
    layout = [
        [SG.Listbox(ps, size=(maiorString, 10), key='PresetSelected')],
        [SG.Button('Atribuir Preset')]
    ]
    return SG.Window('Preset', layout=layout, finalize=True)


Presets = []
caminho = os.path.abspath(os.getcwd())
for diretorio, subpastas, arquivos in os.walk(f'{caminho}\\presets'):
    Presets = arquivos
    maiorString = 0
    for arquivo in arquivos:
        if len(arquivo) > maiorString:
            maiorString = len(arquivo)

x = 1
janela_2, janela_x = None, janela_nova(x)
values = {}
rodar_arquivo = salvar_preset = atribuir_preset = None

# Interação com interface gráfica
while True:
    window, event, value_inst = SG.read_all_windows()
    if event == SG.WIN_CLOSED:
        break
    if window == janela_x and event == 'Selecionar Preset':
        janela_x.hide()
        janela_2 = janela_preset(Presets)
    if window == janela_2 and event == 'Atribuir Preset':
        atribuir_preset = True
        pres = str(value_inst["PresetSelected"])[2:-2]
        values |= value_inst
        janela_2.hide()
        janela_x = janela_nova(x, pres)
    if window == janela_x and event == 'Gerar Arquivo':
        values |= value_inst
        rodar_arquivo = True
        break
    if window == janela_x and event == 'Salvar Preset':
        values |= value_inst
        salvar_preset = True
        break
    if window == janela_x and event == 'Nova requisição' and atribuir_preset is not True:
        x += 1
        values |= value_inst
        janela_x.hide()
        janela_x = janela_nova(x)
    if window == janela_x and event == 'Nova requisição' and atribuir_preset is True:
        x += 1
        values |= value_inst
        janela_x.hide()
        janela_x = janela_nova(x, pres)


# Cria Preset
if salvar_preset is True:
    try:
        os.mkdir(f'{caminho}\\presets')
        preset = open(f'{caminho}\\presets\\PRESET {values["presetName"]}.txt', 'w')
    except FileExistsError:
        preset = open(f'{caminho}\\presets\\PRESET {values["presetName"]}.txt', 'w')
    for key in values:
        escrita = f'{key}: {values[key]}'
        if "presetName" == str(key):
            preset.write(f'{escrita}')
        else:
            preset.write(f'{escrita[3::]}\n')
    preset.close()

# Cria documento
if rodar_arquivo is True:
    documento = Document()
    secoes = documento.sections

    for secao in secoes:
        secao.top_margin = Cm(1.27)
        secao.bottom_margin = Cm(1.27)
        secao.left_margin = Cm(1.27)
        secao.right_margin = Cm(1.27)
        secao.page_height = Cm(29.7)
        secao.page_width = Cm(21.0)
        secao.line_space = Cm(0)

    styles = documento.styles
    p = styles.add_style('p', WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = 'Times New Roman'
    p.font.size = Pt(4)

    Verm = styles.add_style('Verm', WD_STYLE_TYPE.PARAGRAPH)
    Verm.font.name = 'Arial Narrow'
    Verm.font.size = Pt(12)
    Verm.font.color.rgb = RGBColor(255, 0, 0)

    PretN = styles.add_style('PretN', WD_STYLE_TYPE.PARAGRAPH)
    PretN.font.name = 'Arial Narrow'
    PretN.font.size = Pt(12)
    PretN.font.bold = False

    BaseTabela = styles.add_style('BaseTabela', WD_STYLE_TYPE.TABLE)
    BaseTabela.base_style = styles.add_style('TableGrid', WD_STYLE_TYPE.TABLE)
    BaseTabela.font.name = 'Arial Narrow'
    BaseTabela.font.size = Pt(10)
    BaseTabela.font.bold = True

    for l in range(x):
        l += 1
        # Cria a base das tabelas
        # Tabela 1
        cabecalho = documento.add_table(rows=1, cols=4, style='BaseTabela')
        cabecalho.autofit = False

        # \/ Escritas \/
        if True:
            Cab1 = cabecalho.cell(row_idx=0, col_idx=0)
            paraCab1 = Cab1.paragraphs[0]
            paraCab1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runCab1 = paraCab1.add_run()
            runCab1.add_picture('logo.png', width=Cm(3.5), height=Cm(1))
            runCab1.font.name = 'Impact'
            runCab1.font.size = Pt(11)
            runCab1.font.bold = False
            Cab1.width = Cm(4.2)

            Cab2 = cabecalho.cell(row_idx=0, col_idx=1)
            paraCab2 = Cab2.paragraphs[0]
            paraCab2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runCab2 = paraCab2.add_run('Requisição de Trabalho')
            runCab2.font.size = Pt(26)
            Cab2.width = Cm(9.615)

            Cab3 = cabecalho.cell(row_idx=0, col_idx=2)
            Cab3.text = f'Ano:'
            paraCab3 = Cab3.add_paragraph(f'', 'Verm')
            paraCab3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runCab3 = paraCab3.add_run(values[f"{l}: ano"])
            runCab3.font.size = Pt(14)
            Cab3.width = Cm(2.18)

            Cab4 = cabecalho.cell(row_idx=0, col_idx=3)
            Cab4.text = f'Requisição:'
            paraCab4 = Cab4.add_paragraph(f'', 'Verm')
            paraCab4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runCab4 = paraCab4.add_run(values[f"{l}: requisicao"])
            runCab4.font.size = Pt(14)
            Cab4.width = Cm(2.85)

        paragrafo = documento.add_paragraph('', style='p')
        paragrafo.paragraph_format.space_after = Cm(0)

        # Tabela 2
        professores = documento.add_table(rows=2, cols=5, style='BaseTabela')

        # \/ Junta celulas
        if True:
            celula_para_juntar_DS1 = professores.cell(row_idx=0, col_idx=1)
            celula_para_juntar_DS2 = professores.cell(row_idx=0, col_idx=3)
            celula_Disciplina_Setor = celula_para_juntar_DS1.merge(celula_para_juntar_DS2)

            celula_para_juntar_Prof1 = professores.cell(row_idx=0, col_idx=4)
            celula_para_juntar_Prof2 = professores.cell(row_idx=1, col_idx=4)
            celula_Professores = celula_para_juntar_Prof1.merge(celula_para_juntar_Prof2)

        # \/ Escritas \/
        if True:
            Prof1 = professores.cell(row_idx=0, col_idx=0)
            Prof1.text = f'Coordenador:'
            paraProf1 = Prof1.add_paragraph(f'', 'Verm')
            paraProf1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runProf1 = paraProf1.add_run(values[f"{l}: coordenador"])
            runProf1.font.size = Pt(10)
            Prof1.width = Cm(4.85)

            Prof2 = professores.cell(row_idx=0, col_idx=1)
            Prof2.text = f'Disciplina/Setor:'
            paraProf2 = Prof2.add_paragraph(f'', 'Verm')
            paraProf2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runProf2 = paraProf2.add_run(values[f"{l}: DS"])

            Prof3 = professores.cell(row_idx=0, col_idx=4)
            Prof3.text = f'Professores:'
            paraProf3 = Prof3.add_paragraph('', 'PretN')
            paraProf3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runProf3 = paraProf3.add_run(values[f"{l}: professores"])
            runProf3.font.bold = True
            Prof3.width = Cm(7.9)

            Prof4 = professores.cell(row_idx=1, col_idx=0)
            Prof4.text = f'Supervisor'
            paraProf4 = Prof4.add_paragraph(f'', 'Verm')
            paraProf4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runProf4 = paraProf4.add_run(values[f"{l}: supervisor"])
            runProf4.font.size = Pt(10)

            Prof5 = professores.cell(row_idx=1, col_idx=1)
            Prof5.text = f'Sedes:'
            Prof5.width = Cm(1.16)

            Prof6 = professores.cell(row_idx=1, col_idx=2)
            Prof6.text = f'Aldeota'
            paraProf6 = Prof6.add_paragraph()
            runProf6 = paraProf6.add_run()
            add_caixa(runProf6, values[f"{l}: EBS"], 'EBS')
            runProf6.font.size = Pt(12)
            Prof6.width = Cm(1.85)

            Prof7 = professores.cell(row_idx=1, col_idx=3)
            Prof7.text = f'Centro'
            paraProf7 = Prof7.add_paragraph()
            runProf7 = paraProf7.add_run()
            add_caixa(runProf7, values[f"{l}: EGS"], 'EGS')
            runProf7.add_text('  ')
            add_caixa(runProf7, values[f"{l}: NGS"], 'NGS')
            runProf7.font.size = Pt(12)
            Prof2.width = Prof7.width = Cm(3)

        paragrafo = documento.add_paragraph('', style='p')
        paragrafo.paragraph_format.space_after = Cm(0)

        # Tabela 3
        turmas = documento.add_table(rows=2, cols=4, style='BaseTabela')
        # \/ Junta celulas
        if True:
            celula_para_juntar_Etapa1 = turmas.cell(row_idx=0, col_idx=1)
            celula_para_juntar_Etapa2 = turmas.cell(row_idx=1, col_idx=1)
            celula_Etapa = celula_para_juntar_Etapa1.merge(celula_para_juntar_Etapa2)

            celula_para_juntar_Ensino1 = turmas.cell(row_idx=0, col_idx=2)
            celula_para_juntar_Ensino2 = turmas.cell(row_idx=1, col_idx=2)
            celula_Ensino = celula_para_juntar_Ensino1.merge(celula_para_juntar_Ensino2)

        # \/ Escritas \/
        if True:
            Turm1 = turmas.cell(row_idx=0, col_idx=0)
            Turm1.text = f'Turno:'
            paraTurm1 = Turm1.add_paragraph()
            runTurm1 = paraTurm1.add_run()
            add_caixa(runTurm1, values[f"{l}: M"], 'M')
            runTurm1.add_text('         ')
            add_caixa(runTurm1, values[f"{l}: T"], 'T')
            runTurm1.add_text('         ')
            add_caixa(runTurm1, values[f"{l}: N"], 'N')
            runTurm1.font.bold = False

            Turm2 = turmas.cell(row_idx=0, col_idx=1)
            Turm2.text = f'Etapa:'
            paraTurm2_1 = Turm2.add_paragraph()
            Turm2.add_paragraph()
            paraTurm2_2 = Turm2.add_paragraph()
            Turm2.add_paragraph()
            paraTurm2_3 = Turm2.add_paragraph()
            Turm2.add_paragraph()
            paraTurm2_4 = Turm2.add_paragraph()
            runTurm2_1 = paraTurm2_1.add_run()
            runTurm2_2 = paraTurm2_2.add_run()
            runTurm2_3 = paraTurm2_3.add_run()
            runTurm2_4 = paraTurm2_4.add_run()
            add_caixa(runTurm2_1, values[f"{l}: 1et"], '1ª')
            runTurm2_1.add_text('           ')
            add_caixa(runTurm2_1, values[f"{l}: 4et"], '4ª')
            add_caixa(runTurm2_2, values[f"{l}: 2et"], '2ª')
            runTurm2_2.add_text('           ')
            add_caixa(runTurm2_2, values[f"{l}: 2vg"], '2ªVG')
            add_caixa(runTurm2_3, values[f"{l}: 1vg"], '1ªVG')
            runTurm2_3.add_text('      ')
            add_caixa(runTurm2_3, values[f"{l}: rec"], 'Rec.')
            add_caixa(runTurm2_4, values[f"{l}: 3et"], '3ª')
            runTurm2_4.add_text('           ')
            add_caixa(runTurm2_4, values[f"{l}: aff"], 'AF')
            runTurm2_1.font.bold = False
            runTurm2_2.font.bold = False
            runTurm2_3.font.bold = False
            runTurm2_4.font.bold = False
            Turm2.width = Cm(3.33)

            Turm3 = turmas.cell(row_idx=0, col_idx=2)
            Turm3.text = f'Ensino:'
            paraTurm3_1 = Turm3.add_paragraph()
            Turm3.add_paragraph()
            paraTurm3_2 = Turm3.add_paragraph()
            Turm3.add_paragraph()
            paraTurm3_3 = Turm3.add_paragraph()
            Turm3.add_paragraph()
            paraTurm3_4 = Turm3.add_paragraph()
            runTurm3_1 = paraTurm3_1.add_run()
            runTurm3_2 = paraTurm3_2.add_run()
            runTurm3_3 = paraTurm3_3.add_run()
            runTurm3_4 = paraTurm3_4.add_run()
            add_caixa(runTurm3_1, values[f"{l}: infan"], 'Infantil')
            add_caixa(runTurm3_2, values[f"{l}: fund1"], 'Fund. I')
            add_caixa(runTurm3_3, values[f"{l}: fund2"], 'Fund. II')
            add_caixa(runTurm3_4, values[f"{l}: medio"], 'Médio')
            runTurm3_1.font.bold = False
            runTurm3_2.font.bold = False
            runTurm3_3.font.bold = False
            runTurm3_4.font.bold = False
            Turm3.width = Cm(2.2)

            Turm4 = turmas.cell(row_idx=0, col_idx=3)
            Turm4.text = f'Ano:'
            paraTurm4_1 = Turm4.add_paragraph()
            paraTurm4_2 = Turm4.add_paragraph()
            runTurm4_1 = paraTurm4_1.add_run()
            runTurm4_2 = paraTurm4_2.add_run()
            add_caixa(runTurm4_1, values[f"{l}: 1ano"], '1º')
            runTurm4_1.add_text('         ')
            add_caixa(runTurm4_1, values[f"{l}: 2ano"], '2º')
            runTurm4_1.add_text('         ')
            add_caixa(runTurm4_1, values[f"{l}: 3ano"], '3º')
            runTurm4_1.add_text('         ')
            add_caixa(runTurm4_1, values[f"{l}: 4ano"], '4º')
            runTurm4_1.add_text('         ')
            add_caixa(runTurm4_1, values[f"{l}: 5ano"], '5º')
            runTurm4_1.add_text('         ')
            add_caixa(runTurm4_1, values[f"{l}: 6ano"], '6º')
            runTurm4_1.add_text('         ')
            add_caixa(runTurm4_1, values[f"{l}: 7ano"], '7º')
            add_caixa(runTurm4_2, values[f"{l}: 8ano"], '8º')
            runTurm4_2.add_text('         ')
            add_caixa(runTurm4_2, values[f"{l}: 9ano"], '9º')
            runTurm4_2.add_text('         ')
            add_caixa(runTurm4_2, values[f"{l}: ext."], 'Ext.')
            runTurm4_2.add_text('      ')
            add_caixa(runTurm4_2, values[f"{l}: out"], '')
            runTurm4_2.add_text(values[f"{l}: serie"])
            runTurm4_2.add_text(((25 - len(values[f"{l}: serie"])) * '_'))
            runTurm4_1.font.bold = False
            runTurm4_2.font.bold = False

            Turm5 = turmas.cell(row_idx=1, col_idx=0)
            Turm5.text = f'Turma(s):'
            paraTurm5 = Turm5.add_paragraph()
            paraTurm5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runTurm5 = paraTurm5.add_run(values[f"{l}: turmas"])
            runTurm5.font.size = Pt(16)
            runTurm5.font.color.rgb = RGBColor(255, 0, 0)
            Turm1.width = Turm5.width = Cm(3.93)

            Turm6 = turmas.cell(row_idx=1, col_idx=3)
            Turm6.text = f'Natureza do Trabalho'
            paraTurm6_1 = Turm6.add_paragraph()
            paraTurm6_1_1 = Turm6.add_paragraph('', style='p')
            paraTurm6_2 = Turm6.add_paragraph()
            paraTurm6_2_1 = Turm6.add_paragraph('', style='p')
            paraTurm6_3 = Turm6.add_paragraph()
            runTurm6_1 = paraTurm6_1.add_run()
            runTurm6_1_1 = paraTurm6_1_1.add_run('')
            runTurm6_2 = paraTurm6_2.add_run()
            runTurm6_3 = paraTurm6_3.add_run()
            add_caixa(runTurm6_1, values[f"{l}: NP"], 'NP')
            runTurm6_1.add_text('                            ')
            add_caixa(runTurm6_1, values[f"{l}: VG"], 'VG')
            runTurm6_1.add_text('                          ')
            add_caixa(runTurm6_1, values[f"{l}: MT"], 'Miniteste')
            add_caixa(runTurm6_2, values[f"{l}: AP"], 'AP')
            runTurm6_2.add_text('                            ')
            add_caixa(runTurm6_2, values[f"{l}: SIM"], 'Simulados')
            runTurm6_2.add_text('              ')
            add_caixa(runTurm6_2, values[f"{l}: ENEM"], 'SIMULADO ENEM')
            add_caixa(runTurm6_3, values[f"{l}: TD"], 'TD/Tarefas')
            runTurm6_3.add_text('               ')
            add_caixa(runTurm6_3, values[f"{l}: PA"], 'Pautas de Avaliação')
            runTurm6_1.font.bold = False
            runTurm6_2.font.bold = False
            runTurm6_3.font.bold = False
            Turm4.width = Turm6.width = Cm(10)

        paragrafo = documento.add_paragraph('', style='p')
        paragrafo.paragraph_format.space_after = Cm(0)

        # Tabela 4
        observacoes = documento.add_table(rows=2, cols=3, style='BaseTabela')
        # \/ Junta celulas
        if True:
            celula_para_juntar_Observacoes1 = observacoes.cell(row_idx=0, col_idx=0)
            celula_para_juntar_Observacoes2 = observacoes.cell(row_idx=0, col_idx=2)
            celula_Observacoes = celula_para_juntar_Observacoes1.merge(celula_para_juntar_Observacoes2)

        # \/ Escritas \/
        if True:
            Obs1 = observacoes.cell(row_idx=0, col_idx=0)
            paraObs1 = Obs1.paragraphs[0]
            paraObs1.add_run(f'Observação: {values[f"{l}: observacoes"]}')

            Obs2 = observacoes.cell(row_idx=1, col_idx=0)
            Obs2.text = f'Envio:'
            paraObs2 = Obs2.add_paragraph()
            runObs2 = paraObs2.add_run(values[f"{l}: envio"])
            runObs2.font.size = Pt(12)
            runObs2.font.color.rgb = RGBColor(255, 0, 0)

            Obs3 = observacoes.cell(row_idx=1, col_idx=1)
            Obs3.text = f'Aplicação:'
            paraObs3 = Obs3.add_paragraph()
            runObs3 = paraObs3.add_run(values[f"{l}: aplicacao"])
            runObs3.font.size = Pt(12)
            runObs3.font.color.rgb = RGBColor(255, 0, 0)
            Obs2.width = Obs3.width = Cm(5.11)

            Obs4 = observacoes.cell(row_idx=1, col_idx=2)
            Obs4.text = f'SOP Correção:'
            paraObs4 = Obs4.add_paragraph()
            runObs4 = paraObs4.add_run()
            add_caixa(runObs4, values[f"{l}: Sim"], 'Sim')
            runObs4.add_text('         ')
            add_caixa(runObs4, values[f"{l}: Nao"], 'Não')
            runObs4.add_text('         ')
            runObs4.add_text('Questões: ')
            runObs4.add_text(values[f"{l}: SOP"])
            runObs4.add_text(((40 - len(values[f"{l}: SOP"])) * '_'))
            runObs4.font.bold = False
            Obs4.width = Cm(18)

        paragrafo = documento.add_paragraph('', style='p')
        paragrafo.paragraph_format.space_after = Cm(0)

        # Tabela 5
        uso_da_grafica = documento.add_table(rows=8, cols=9, style='BaseTabela')
        # \/ Junta celulas
        if True:
            celula_para_juntar_PUG1 = uso_da_grafica.cell(row_idx=0, col_idx=0)
            celula_para_juntar_PUG2 = uso_da_grafica.cell(row_idx=0, col_idx=8)
            celula_para_uso_grafica = celula_para_juntar_PUG1.merge(celula_para_juntar_PUG2)

            celula_para_juntar_Hora1 = uso_da_grafica.cell(row_idx=1, col_idx=1)
            celula_para_juntar_Hora2 = uso_da_grafica.cell(row_idx=1, col_idx=2)
            celula_Hora = celula_para_juntar_Hora1.merge(celula_para_juntar_Hora2)

            celula_para_juntar_Qtd1 = uso_da_grafica.cell(row_idx=1, col_idx=3)
            celula_para_juntar_Qtd2 = uso_da_grafica.cell(row_idx=1, col_idx=6)
            celula_Qtd = celula_para_juntar_Qtd1.merge(celula_para_juntar_Qtd2)

            celula_para_juntar_Arte1 = uso_da_grafica.cell(row_idx=1, col_idx=7)
            celula_para_juntar_Arte2 = uso_da_grafica.cell(row_idx=1, col_idx=8)
            celula_Arte = celula_para_juntar_Arte1.merge(celula_para_juntar_Arte2)

            # Junta celulas de termino e Visto
            for linha in range(6):
                linha += 2
                celula_para_juntar_Termino1 = uso_da_grafica.cell(row_idx=linha, col_idx=2)
                celula_para_juntar_Termino2 = uso_da_grafica.cell(row_idx=linha, col_idx=3)
                celula_Termino = celula_para_juntar_Termino1.merge(celula_para_juntar_Termino2)

                celula_para_juntar_Visto1 = uso_da_grafica.cell(row_idx=linha, col_idx=6)
                celula_para_juntar_Visto2 = uso_da_grafica.cell(row_idx=linha, col_idx=7)
                celula_Visto = celula_para_juntar_Visto1.merge(celula_para_juntar_Visto2)

            celula_para_juntar_DV1 = uso_da_grafica.cell(row_idx=1, col_idx=3)
            celula_para_juntar_DV2 = uso_da_grafica.cell(row_idx=1, col_idx=6)
            celula_data_e_visto = celula_para_juntar_DV1.merge(celula_para_juntar_DV2)

        # \/ Escritas \/
        if True:
            paraUG0 = uso_da_grafica.cell(row_idx=0, col_idx=0).paragraphs[0]
            RUG0 = paraUG0.add_run('PARA USO DA GRÁFICA')
            uso_da_grafica.cell(row_idx=0, col_idx=0).paragraphs[0].paragraph_format.alignment = \
                WD_PARAGRAPH_ALIGNMENT.CENTER
            RUG0.font.size = Pt(12)
            uso_da_grafica.cell(row_idx=0, col_idx=0).width = Cm(20)

            RUG1 = uso_da_grafica.cell(row_idx=1, col_idx=0).paragraphs[0].add_run('Recebimento: ____/_____/_____')
            RUG1.font.bold = False
            uso_da_grafica.cell(row_idx=1, col_idx=0).width = Cm(4.75)

            RUG1 = uso_da_grafica.cell(row_idx=1, col_idx=1).paragraphs[0].add_run('Hora: _____ : _____')
            RUG1.font.bold = False
            uso_da_grafica.cell(row_idx=1, col_idx=1).width = Cm(3.25)

            RUG1 = uso_da_grafica.cell(row_idx=1, col_idx=3).paragraphs[0].add_run('Qtde. Solicitada: ______________')
            RUG1.font.bold = False
            uso_da_grafica.cell(row_idx=1, col_idx=3).width = Cm(5.5)
            uso_da_grafica.rows[1].height = Cm(0.35)

            RUG1 = uso_da_grafica.cell(row_idx=1, col_idx=7).paragraphs[0].add_run('Arte Fornecida? (S/N) __________')
            RUG1.font.bold = False
            uso_da_grafica.cell(row_idx=2, col_idx=0).text = 'Etapas de Elaboração'

            paraUG2 = uso_da_grafica.cell(row_idx=2, col_idx=1).paragraphs[0]
            RUG2 = paraUG2.add_run('Início')
            paraUG2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            uso_da_grafica.cell(row_idx=2, col_idx=1).width = \
                uso_da_grafica.cell(row_idx=3, col_idx=1).width = \
                uso_da_grafica.cell(row_idx=4, col_idx=1).width = \
                uso_da_grafica.cell(row_idx=5, col_idx=1).width = \
                uso_da_grafica.cell(row_idx=6, col_idx=1).width = \
                uso_da_grafica.cell(row_idx=7, col_idx=1).width = Cm(1.90)

            paraUG3 = uso_da_grafica.cell(row_idx=2, col_idx=2).paragraphs[0]
            RUG3 = paraUG3.add_run('Término')
            paraUG3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            uso_da_grafica.cell(row_idx=2, col_idx=2).width = \
                uso_da_grafica.cell(row_idx=3, col_idx=2).width = \
                uso_da_grafica.cell(row_idx=4, col_idx=2).width = \
                uso_da_grafica.cell(row_idx=5, col_idx=2).width = \
                uso_da_grafica.cell(row_idx=6, col_idx=2).width = \
                uso_da_grafica.cell(row_idx=7, col_idx=2).width = Cm(1.77)

            paraUG4 = uso_da_grafica.cell(row_idx=2, col_idx=4).paragraphs[0]
            RUG4 = paraUG4.add_run('Qtde.\nHoras')
            paraUG4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            paraUG5 = uso_da_grafica.cell(row_idx=2, col_idx=5).paragraphs[0]
            RUG5 = paraUG5.add_run('Hr. Extras')
            paraUG5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            paraUG6 = uso_da_grafica.cell(row_idx=2, col_idx=6).paragraphs[0]
            RUG6 = paraUG6.add_run('Visto')
            paraUG6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            uso_da_grafica.cell(row_idx=2, col_idx=6).width = \
                uso_da_grafica.cell(row_idx=3, col_idx=6).width = \
                uso_da_grafica.cell(row_idx=4, col_idx=6).width = \
                uso_da_grafica.cell(row_idx=5, col_idx=6).width = \
                uso_da_grafica.cell(row_idx=6, col_idx=6).width = \
                uso_da_grafica.cell(row_idx=7, col_idx=6).width = Cm(2)

            uso_da_grafica.cell(row_idx=2, col_idx=8).text = 'Data e Visto da Impressão'
            RUG1 = uso_da_grafica.cell(row_idx=3, col_idx=0).paragraphs[0].add_run('Desenho')
            RUG1.font.bold = False
            RUG1 = uso_da_grafica.cell(row_idx=4, col_idx=0).paragraphs[0].add_run('Digitação')
            RUG1.font.bold = False
            RUG1 = uso_da_grafica.cell(row_idx=5, col_idx=0).paragraphs[0].add_run('Revisão')
            RUG1.font.bold = False
            RUG1 = uso_da_grafica.cell(row_idx=6, col_idx=0).paragraphs[0].add_run('Produção')
            RUG1.font.bold = False
            RUG1 = uso_da_grafica.cell(row_idx=7, col_idx=0).paragraphs[0].add_run('Acabamento')
            RUG1.font.bold = False

        paragrafo = documento.add_paragraph('', style='p')
        paragrafo.paragraph_format.space_after = Cm(0)
        paragrafo = documento.add_paragraph('', style='p')
        paragrafo.paragraph_format.space_after = Cm(0)
        if l % 2 == 1:
            paragrafo = documento.add_paragraph('', style='p')
            paragrafo.paragraph_format.space_after = Cm(0)
            paragrafo = documento.add_paragraph('', style='p')
            paragrafo.paragraph_format.space_after = Cm(0)
            paragrafo = documento.add_paragraph('', style='p')
            paragrafo.paragraph_format.space_after = Cm(0)

    caminho = os.path.abspath(os.getcwd())
    todas_as_requisicoes = ''
    for d in range(x):
        d += 1
        if d < x:
            if d == x - 1:
                todas_as_requisicoes += f'{values[f"{d}: requisicao"]} e '
            else:
                todas_as_requisicoes += f'{values[f"{d}: requisicao"]}, '
        else:
            todas_as_requisicoes += f'{values[f"{d}: requisicao"]}'
    try:
        os.mkdir(f'{caminho}\\Retorno')
        documento.save(f'{caminho}\\Retorno\\REQ DE TAREFAS {todas_as_requisicoes}.docx')
    except FileExistsError:
        documento.save(f'{caminho}\\Retorno\\REQ DE TAREFAS {todas_as_requisicoes}.docx')
